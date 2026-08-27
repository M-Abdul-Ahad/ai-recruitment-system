import re
import docx
import PyPDF2
from resumes.data.skills import SKILLS

# Date pattern to detect job dates/durations in experience sections
DATE_PATTERN = (
    r'(?:\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|'
    r'aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?|\d{1,2}[/-])\s*)?'
    r'\b(?:19|20)\d{2}\b'
    r'(?:\s*(?:[-–]|to|\s)\s*'
    r'(?:\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|'
    r'aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?|\d{1,2}[/-])\s*)?'
    r'(?:\b(?:19|20)\d{2}\b|present|current|now))?'
)


def extract_text_from_pdf(file_input):
    text = ""
    if isinstance(file_input, (str, bytes)):
        with open(file_input, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    else:
        file_input.seek(0)
        reader = PyPDF2.PdfReader(file_input)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_text_from_docx(file_input):
    if isinstance(file_input, (str, bytes)):
        doc = docx.Document(file_input)
    else:
        file_input.seek(0)
        doc = docx.Document(file_input)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def clean_resume_text(text):
    text = text.lower()
    text = re.sub(r'\n+', '\n', text)              # keep structure
    text = re.sub(r'[^\w\s\n]', '', text)          # keep \n
    text = re.sub(r'[ \t]+', ' ', text)            # only spaces
    return text.strip()


def extract_skills(cleaned_text):
    found_skills = []

    for skill in SKILLS:
        if skill in cleaned_text:
            found_skills.append(skill)

    return list(set(found_skills))


def extract_education_section(cleaned_text):
    lines = cleaned_text.split('\n')
    education_lines = []
    capture = False

    start_keywords = [
        "education", "academic", "academics",
        "qualification", "qualifications",
        "educational background", "academic background"
    ]

    stop_keywords = [
        "experience", "work", "skills",
        "projects", "certifications", "internships"
    ]

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if any(k in line for k in start_keywords):
            capture = True
            continue

        if capture and any(k in line for k in stop_keywords):
            break

        if capture:
            education_lines.append(line)

    return education_lines


def extract_year(text):
    match = re.search(r'(19|20)\d{2}', text)
    return match.group() if match else ""


def parse_education(education_lines):
    education_data = []

    for line in education_lines:
        degree = ""
        institution = ""
        year = ""

        # Degree detection
        for keyword in [
            "bachelor", "bsc", "bs",
            "master", "msc", "ms",
            "phd", "diploma"
        ]:
            if keyword in line:
                degree = keyword.title()
                break

        # Institution detection
        if any(word in line for word in ["university", "college", "institute"]):
            institution = re.sub(r'(19|20)\d{2}', '', line).title().strip()

        # Year detection
        year_match = re.search(r'(19|20)\d{2}', line)
        if year_match:
            year = year_match.group()

        if degree or institution:
            education_data.append({
                "degree": degree,
                "institution": institution,
                "year": year
            })

    return education_data


def extract_experience_section(cleaned_text):
    lines = cleaned_text.split('\n')
    experience_lines = []
    capture = False

    start_keywords = [
        "experience", "work experience",
        "employment", "professional experience",
        "internship", "industrial training"
    ]

    stop_keywords = [
        "education", "skills",
        "projects", "certifications", "awards"
    ]

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if any(k in line for k in start_keywords):
            capture = True
            continue

        if capture and any(k in line for k in stop_keywords):
            break

        if capture:
            experience_lines.append(line)

    return experience_lines


def split_experience_blocks(experience_lines):
    blocks = []
    current_block = []

    for line in experience_lines:
        if re.search(DATE_PATTERN, line, re.IGNORECASE):
            if current_block:
                blocks.append(current_block)
                current_block = []
        current_block.append(line)

    if current_block:
        blocks.append(current_block)

    return blocks


def parse_experience(blocks):
    experience_data = []

    for block in blocks:
        job_title = ""
        company = ""
        duration = ""
        description_lines = []

        # First line usually contains title + company + duration
        header = block[0]

        # Duration
        duration_match = re.search(
            r'(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|[0-9]{1,2}/)\s*)?'
            r'(?:19|20)\d{2}\s*[-–\s|to]*\s*'
            r'(?:(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|[0-9]{1,2}/)\s*)?'
            r'(?:(?:19|20)\d{2}|present|current|now)',
            header,
            re.IGNORECASE
        )
        if duration_match:
            duration = duration_match.group()

        # Split title and company
        if " at " in header:
            job_title, company = header.split(" at ", 1)
        elif " - " in header:
            job_title, company = header.split(" - ", 1)
        else:
            job_title = header
            company = ""

        # Remaining lines → description
        if len(block) > 1:
            description_lines = block[1:]

        experience_data.append({
            "job_title": job_title,
            "company": company,
            "duration": duration,
            "description": " ".join(description_lines)
        })

    return experience_data
